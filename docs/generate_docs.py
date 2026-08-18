import os
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell margins (padding) in dxa (1/20 of a pt)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def set_run_font(run, name='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_p(doc, text="", bold=False, italic=False, size=12, alignment=None, space_after=6, space_before=0, line_spacing=1.15):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        p.alignment = alignment
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p

def add_run(p, text, bold=False, italic=False, size=12):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return run

def modify_mtz_docx():
    print("Generating Відомості_МТЗ_ЗОІППО.docx from template in landscape orientation...")
    template_path = 'lms-portal/docs/Відомості про МТЗ_шаблон.docx'
    doc_path = 'lms-portal/docs/Відомості_МТЗ_ЗОІППО.docx'
    doc = Document(template_path)
    
    def set_cell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=10):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold

    # 1. Fill Header P[2] and P[3]
    doc.paragraphs[2].text = ''
    r1 = doc.paragraphs[2].add_run('Заявник ')
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)

    r2 = doc.paragraphs[2].add_run('Комунальний заклад «Запорізький обласний інститут післядипломної педагогічної освіти» Запорізької обласної ради, код ЄДРПОУ 02136146. Структурний підрозділ: Організаційно-методичний центр\n')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12); r2.bold = True

    r3 = doc.paragraphs[2].add_run('(найменування, код згідно з ЄДРПОУ, у разі коли статус кваліфікаційного центру має намір набути структурний або відокремлений підрозділ заявника, - також найменування такого підрозділу)')
    r3.font.name = 'Times New Roman'; r3.font.size = Pt(10); r3.italic = True

    doc.paragraphs[3].text = ''
    r4 = doc.paragraphs[3].add_run('Назва професійної кваліфікації “')
    r4.font.name = 'Times New Roman'; r4.font.size = Pt(12)

    r5 = doc.paragraphs[3].add_run('Фахівець із супроводу ветеранів війни та демобілізованих осіб')
    r5.font.name = 'Times New Roman'; r5.font.size = Pt(12); r5.bold = True

    r6 = doc.paragraphs[3].add_run('”')
    r6.font.name = 'Times New Roman'; r6.font.size = Pt(12)

    # 2. Modify Table 0 (7 columns)
    t0 = doc.tables[0]
    row0 = t0.add_row()
    t0_data = [
        '1',
        'Банк тестових питань для оцінювання результатів навчання за професійною кваліфікацією «Фахівець із супроводу ветеранів війни та демобілізованих осіб»',
        '100 питань, 2 варіанти (Варіант 1: 50 питань; Варіант 2: 50 питань)',
        'Тестові завдання складаються з: завдання одиничного вибору (з однією правильною відповіддю із чотирьох запропонованих варіантів).',
        '02.02.2026 - 25.03.2026; слухачі курсів підвищення кваліфікації; категорія - соціальні працівники, соціальні педагоги та практичні психологи; кількість учасників — 136; апробацію проведено в електронній системі на базі Moodle в КЗ «Запорізький обласний інститут післядипломної педагогічної освіти» ЗОР.',
        '02.02.2026 - 25.03.2026 проведено представниками КЗ «Запорізький обласний інститут післядипломної педагогічної освіти» ЗОР: проректором з навчально-методичної роботи Черніковою Л.А., завідувачем обласного науково-методичного центру інформатизації освіти Здоровцем О.Ф., завідувачем навчально-методичного центру Задорожкіною Я.С. за результатами статистичного аналізу та апробації (коефіцієнт внутрішньої узгодженості перевищує 84%, розроблені тестові та практичні завдання є валідними та задовільними).',
        '1) Оформлено інструкцію про ознайомлення співробітників із захистом відомостей про тестові завдання та нормами юридичної відповідальності за порушення. 2) Обмежено фізичний доступ сторонніх осіб до тестування.'
    ]
    for idx, text in enumerate(t0_data):
        align = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
        set_cell(row0.cells[idx], text, align=align, size=10)

    # 3. Modify Table 1 (5 columns)
    t1 = doc.tables[1]
    row1 = t1.add_row()
    t1_data = [
        '1',
        'Пакет практичних кваліфікаційних завдань (професійні кейси)',
        '1 (10 кейсів)',
        'Оцінювання практичних умінь та навичок кандидатів, здатності аналізувати складні ситуації супроводу, приймати рішення у кризових станах, дотримуватися етичних норм та законодавства про захист персональних даних.',
        '02.02.2026 - 25.03.2026; слухачі курсів підвищення кваліфікації; категорія - соціальні працівники, соціальні педагоги та практичні психологи; кількість учасників — 136; апробацію проведено в електронній системі на базі Moodle в КЗ «Запорізький обласний інститут післядипломної педагогічної освіти» ЗОР.'
    ]
    for idx, text in enumerate(t1_data):
        align = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
        set_cell(row1.cells[idx], text, align=align, size=10)

    # 4. Modify Table 2 (6 columns)
    t2 = doc.tables[2]
    t2_rows = [
        ('1', 'Ноутбук HP 250 (10 од.)', 'Для організації онлайн-оцінювання (тестування), проведення розрахунків, підготовки та збереження звітної інформації.', 'м. Запоріжжя, вул. Незалежної України, 57А, підвальне приміщення літ. А-4', 'КЗ «ЗОІППО» ЗОР', 'Власність закладу, без строку дії (інв. № 11137103)'),
        ('2', 'Веб-камера широкоформатна Logitech C920 Pro HD (1 од.)', 'Відеофіксація процедури оцінювання, ідентифікації кандидатів.', 'м. Запоріжжя, вул. Незалежної України, 57А, підвальне приміщення літ. А-4', 'КЗ «ЗОІППО» ЗОР', 'Власність закладу, без строку дії (інв. № 11139103)'),
        ('3', 'Мультимедійний проектор Acer', 'Демонстрація інструкцій, інструктажів під час оцінювання.', 'м. Запоріжжя, вул. Незалежної України, 57А, підвальне приміщення літ. А-4', 'КЗ «ЗОІППО» ЗОР', 'Власність закладу, без строку дії (інв. № 101490248)'),
        ('4', 'Меблі офісні (столи розкладні - 8 од., стільці м\'які - 24 од.)', 'Для зручного розташування здобувачів та членів комісії під час проведення теоретичного та практичного оцінювання.', 'м. Запоріжжя, вул. Незалежної України, 57А, підвальне приміщення літ. А-4', 'КЗ «ЗОІППО» ЗОР', 'Власність закладу, благодійні, без строку дії (столи STAUNING: інв. № 11139118; стільці м\'які BISTRUP: інв. № 11139113, ULDAI: інв. № 11139114)'),
        ('5', 'Сейф металевий', 'Для надійного та безпечного зберігання конфіденційних матеріалів.', 'м. Запоріжжя, вул. Незалежної України, 57А, підвальне приміщення літ. А-4', 'КЗ «ЗОІППО» ЗОР', 'Власність закладу, без строку дії (інв. № 11136109)'),
        ('6', 'Приміщення Кваліфікаційного центру (загальна площа 658.5 кв. м)', 'Для розташування робочих місць та безпосереднього проведення процедур оцінювання кандидатів в безпечних умовах.', 'м. Запоріжжя, вул. Незалежної України, 57А, підвальне приміщення літ. А-4', 'КЗ «ЗОІППО» ЗОР', 'Власність закладу на підставі свідоцтва про право власності, без строку дії')
    ]

    for row_data in t2_rows:
        row = t2.add_row()
        for idx, text in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(row.cells[idx], text, align=align, size=10)

    # 5. Modify Table 3 (Signature Block)
    t3 = doc.tables[3]
    set_cell(t3.rows[0].cells[0], 'Ректор КЗ «ЗОІППО» ЗОР', align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=11)
    set_cell(t3.rows[0].cells[1], '____________________\n(підпис)', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    set_cell(t3.rows[0].cells[2], 'Едуард ГУГНІН\n(прізвище, ініціали)', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)

    doc.save(doc_path)
    print("Відомості_МТЗ_ЗОІППО.docx generated successfully from template in landscape mode.")

def generate_zayava_docx():
    print("Generating Заява.docx...")
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        
    # Title
    add_p(doc, "ЗАЯВА", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc, "про проведення процедури акредитації", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc, "кваліфікаційного центру", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    # Grid/fields table
    # Columns: 1: Field name, 2: Value
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    fields = [
        ("Заявник", "Комунальний заклад «Запорізький обласний інститут післядипломної педагогічної освіти» Запорізької обласної ради\n(повне найменування юридичної особи)"),
        ("Код згідно з ЄДРПОУ", "02136146"),
        ("Місцезнаходження заявника та його структурного або відокремленого підрозділу, що має намір набути статус кваліфікаційного центру", "69035, м. Запоріжжя, вул. Незалежної України, 57-А"),
        ("Адреса електронної пошти заявника", "osvita@zoippo.zp.ua"),
        ("Номер телефону", "+38 (061) 717-17-72"),
        ("Адреса вебсайта", "http://lms.ele.zp.ua/"),
        ("Назва професійного стандарту, ким і коли затверджений", "Професійний стандарт «Фахівець із супроводу ветеранів війни та демобілізованих осіб», затверджений наказом Міністерства у справах ветеранів України від 10.10.2025 р. № 835"),
        ("Назва професійної кваліфікації", "Фахівець із супроводу ветеранів війни та демобілізованих осіб")
    ]
    
    for f_name, f_val in fields:
        row = table.add_row()
        cell_name, cell_val = row.cells
        
        # Name cell
        cell_name.text = ""
        p_name = cell_name.paragraphs[0]
        p_name.paragraph_format.space_after = Pt(4)
        p_name.paragraph_format.space_before = Pt(4)
        add_run(p_name, f_name, bold=True, size=11)
        
        # Value cell
        cell_val.text = ""
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(4)
        p_val.paragraph_format.space_before = Pt(4)
        add_run(p_val, f_val, size=11)
        
        # Style borders and margins
        for cell in (cell_name, cell_val):
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            set_cell_border(cell,
                            top={'sz': 4, 'val': 'single', 'color': '808080'},
                            bottom={'sz': 4, 'val': 'single', 'color': '808080'},
                            left={'sz': 4, 'val': 'single', 'color': '808080'},
                            right={'sz': 4, 'val': 'single', 'color': '808080'})
            
    # Documents list header
    add_p(doc, "", space_after=12)
    add_p(doc, "Перелік документів, що додаються:", bold=True, size=12, space_before=12, space_after=6)
    
    docs_list = [
        "Копія Положення про Організаційно-методичний центр Комунального закладу «Запорізький обласний інститут післядипломної педагогічної освіти» Запорізької обласної ради на 10 арк. у 1 прим.",
        "Витяг з наказу «Про оптимізацію структури Комунального закладу «Запорізький обласний інститут післядипломної педагогічної освіти» Запорізької обласної ради» від 27.01.2020 № 037 на 1 арк. у 1 прим.",
        "Виписка з Єдиного державного реєстру підприємств та організацій України від 25.10.2024 № 1001031070020008514 на 2 арк. у 1 прим.",
        "Відомості про осіб, яких заявник залучає як оцінювачів на 4 арк. в 1 прим.",
        "Скан-копія трудової книжки Гура Т.Є. на 17 арк. у 1 прим.",
        "Скан-копія трудової книжки Мороко В.В. на 26 арк. у 1 прим.",
        "Скан-копія трудової книжки Мостова Т.О. на 8 арк. у 1 прим.",
        "Порядок присвоєння/підтвердження професійної кваліфікації за професійним стандартом «Фахівець із супроводу ветеранів війни та демобілізованих осіб», назва професійної кваліфікації «Фахівець із супроводу ветеранів війни та демобілізованих осіб» (затверджений рішенням Вченої ради Комунального закладу «Запорізький обласний інститут післядипломної педагогічної освіти» Запорізької обласної ради протокол № 5 від 26.06.2026) на 7 арк. у 1 прим.",
        "Відомості про матеріально-технічне забезпечення процедур оцінювання професійної кваліфікації «Фахівець із супроводу ветеранів війни та демобілізованих осіб» на 4 арк. у 1 прим.",
        "Зразок контрольно-оцінювальних матеріалів для проведення оцінювання результатів навчання на 28 арк. у 1 прим.",
        "Звіт про організацію та проведення апробації тестових та практичних завдань за професійною кваліфікацією «Фахівець із супроводу ветеранів війни та демобілізованих осіб» на 1 арк. у 1 прим.",
        "Відомості про вебсайт організаційно-методичного центру Комунального закладу «Запорізький обласний інститут післядипломної педагогічної освіти» Запорізької обласної ради на 1 арк. у 1 прим.",
        "Статут Комунального закладу «Запорізький обласний інститут післядипломної педагогічної освіти» Запорізької обласної ради на 21 арк. у 1 прим.",
        "Інвентарний опис необоротних активів на 37 арк. у 1 прим.",
        "Документи про право власності на нерухоме майно на 8 арк. у 1 прим.",
        "Документи про освіту Гури Т.Є. на 3 арк. у 1 прим.",
        "Документи про освіту Мороко В.В. на 2 арк. у 1 прим.",
        "Документи про освіту Мостової Т.О. на 3 арк. у 1 прим.",
        "Сертифікати про підвищення кваліфікації Гури Т.Є. на 9 арк. у 1 прим.",
        "Сертифікати про підвищення кваліфікації Мороко В.В. на 3 арк. у 1 прим.",
        "Сертифікати про підвищення кваліфікації Мостової Т.О. на 7 арк. у 1 прим."
    ]
    
    for idx, doc_item in enumerate(docs_list):
        p = add_p(doc, space_after=3)
        p.paragraph_format.left_indent = Cm(0.75)
        add_run(p, f"{idx + 1}. {doc_item}", size=11)
        
    # Signature
    add_p(doc, "", space_after=18)
    p_sig = add_p(doc, space_after=4, space_before=12)
    p_sig.paragraph_format.tab_stops.add_tab_stop(Cm(11.0))
    add_run(p_sig, "Керівник (уповноважена особа)\tЕдуард ГУГНІН", bold=True, size=11)
    
    p_date = add_p(doc, space_after=4, space_before=6)
    p_date.paragraph_format.tab_stops.add_tab_stop(Cm(11.0))
    add_run(p_date, "«16» липня 2026 р.", size=11)
    
    doc.save('lms-portal/docs/Заява.docx')
    print("Заява.docx generated successfully.")

def generate_zvit_docx():
    print("Generating Звіт_з_апробації.docx...")
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        
    # Official Letterhead Header
    p_lh1 = add_p(doc, "КОМУНАЛЬНИЙ ЗАКЛАД «ЗАПОРІЗЬКИЙ ОБЛАСНИЙ ІНСТИТУТ ПІСЛЯДИПЛОМНОЇ ПЕДАГОГІЧНОЇ ОСВІТИ»", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    p_lh2 = add_p(doc, "ЗАПОРІЗЬКОЇ ОБЛАСНОЇ РАДИ", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    p_lh3 = add_p(doc, "вул. Незалежної України, 57-А, м. Запоріжжя, 69035, тел.: (061) 717-17-72", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    p_lh4 = add_p(doc, "E-mail: osvita@zoippo.zp.ua, Код ЄДРПОУ 02136146", size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    
    # Horizontal separator line
    p_sep = add_p(doc, "____________________________________________________________________________________________________", size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    # Ref number and date line
    p_ref = add_p(doc, space_after=18)
    p_ref.paragraph_format.tab_stops.add_tab_stop(Cm(11.0))
    add_run(p_ref, "«27» березня 2026 р.\t№ 01-12/452", bold=True, size=11)
    
    # Report Title
    add_p(doc, "ЗВІТ", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=12)
    add_p(doc, "про організацію та проведення апробації", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc, "контрольно-оцінювальних матеріалів для проведення оцінювання результатів навчання", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc, "за професійною кваліфікацією «Фахівець із супроводу ветеранів війни", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc, "та демобілізованих осіб»", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    # Body Text
    paragraphs_content = [
        "Апробація тестових та практичних завдань проводилась у період з 02 лютого по 25 березня 2026 року в електронній системі на базі Moodle за участі 136 слухачів курсів підвищення кваліфікації — соціальних працівників, соціальних педагогів та практичних психологів, які проходили навчання у КЗ «Запорізький обласний інститут післядипломної педагогічної освіти» ЗОР.",
        "Коефіцієнт внутрішньої узгодженості тестових завдань за результатами апробації перевищує 84%, що свідчить про високу якість, збалансованість та надійність розроблених тестів.",
        "Оцінювання валідності банку тестових завдань за результатами апробації та експертна оцінка змісту були проведені 26 березня 2026 року представниками КЗ «Запорізький обласний інститут післядипломної педагогічної освіти» ЗОР: проректором з навчально-методичної роботи Черніковою Л.А., завідувачем обласного науково-методичного центру інформатизації освіти Здоровцем О.Ф., завідувачем навчально-методичного центру Задорожкіною Я.С.",
        "За результатами статистичного аналізу та експертизи було встановлено, що:\n" +
        "• більшість учасників апробації отримали оцінки у діапазоні від 70 до 85 балів;\n" +
        "• оцінки учасників розподілені рівномірно відповідно до закону нормального розподілу;\n" +
        "• зміст тестових та практичних завдань повністю відповідає вимогам Професійного стандарту «Фахівець із супроводу ветеранів війни та демобілізованих осіб»;\n" +
        "• завдання мають різний рівень складності, що дозволяє диференційовано оцінити знання учасників різного рівня підготовки;\n" +
        "• розроблені практичні завдання (кейси) сприяють перевірці критичних навичок аналізу складних ситуацій супроводу, прийняття професійних рішень та ділової комунікації;\n" +
        "• ліміт часу, виділений на виконання КОМ, є повністю достатнім для понад 95% учасників.",
        "За результатами апробації експертна група робить висновок, що розроблені контрольно-оцінювальні матеріали (тестові та практичні завдання) є валідними, надійними та в повному обсязі придатні для використання у процедурах оцінювання результатів навчання здобувачів професійної кваліфікації «Фахівець із супроводу ветеранів війни та демобілізованих осіб»."
    ]
    
    for p_text in paragraphs_content:
        p = add_p(doc, space_after=6, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        p.paragraph_format.first_line_indent = Cm(1.25)
        # Handle list bullets in fourth paragraph
        if '\n•' in p_text:
            lines = p_text.split('\n')
            p.text = ""
            add_run(p, lines[0], size=11.5)
            for bullet in lines[1:]:
                p_b = add_p(doc, space_after=3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
                p_b.paragraph_format.left_indent = Cm(1.75)
                add_run(p_b, bullet, size=11.5)
        else:
            add_run(p, p_text, size=11.5)
            
    # Signatures
    add_p(doc, "", space_after=24)
    
    signatures = [
        ("Проректор з навчально-методичної роботи", "Людмила ЧЕРНІКОВА"),
        ("Завідувач обласного науково-методичного центру інформатизації освіти", "Олексій ЗДОРОВЕЦЬ"),
        ("Завідувач навчально-методичного центру", "Яна ЗАДОРОЖКІНА")
    ]
    
    for role, name in signatures:
        p_sig = add_p(doc, space_after=4, space_before=4)
        p_sig.paragraph_format.tab_stops.add_tab_stop(Cm(11.0))
        add_run(p_sig, f"{role}\t{name}", bold=True, size=11)
        
    doc.save('lms-portal/docs/Звіт_з_апробації.docx')
    print("Звіт_з_апробації.docx generated successfully.")

def main():
    modify_mtz_docx()
    generate_zayava_docx()
    generate_zvit_docx()
    print("All docs modified and generated successfully!")

if __name__ == '__main__':
    main()
