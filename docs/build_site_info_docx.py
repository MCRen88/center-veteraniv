import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_p(doc, text="", bold=False, italic=False, size=12, alignment=None, space_after=6, space_before=0, line_spacing=1.15):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        p.alignment = alignment
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

def add_run(p, text, bold=False, italic=False, size=12):
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run

def build_site_info_docx():
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)  # standard left margin for binding
        section.right_margin = Cm(1.5)

    # 1. Header
    p_header = add_p(doc, space_after=12, space_before=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p_header, "Кваліфікаційний центр\nКомунального закладу «Запорізький обласний інститут післядипломної\nпедагогічної освіти» Запорізької обласної ради інформує:", bold=True, size=14)

    # Spacer
    add_p(doc, "", space_after=18)

    # 2. Body Text
    p_body = add_p(doc, space_after=12, line_spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p_body.paragraph_format.first_line_indent = Cm(1.25)
    add_run(p_body, "На вебсайті ", size=12)
    add_run(p_body, "http://lms.ele.zp.ua/", bold=True, size=12)
    add_run(p_body, " розміщено інформацію про Кваліфікаційний центр сертифікації фахівців із супроводу ветеранів, контактну інформацію про Кваліфікаційний центр; відповідні професійні стандарти; процедури присвоєння/підтвердження; зразок заяви про присвоєння/підтвердження професійної кваліфікації; перелік документів, які додаються до заяви здобувача.", size=12)

    # Spacer
    add_p(doc, "", space_after=12)

    # 3. Structure
    p_struct_title = add_p(doc, space_after=6, space_before=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p_struct_title.paragraph_format.first_line_indent = Cm(1.25)
    add_run(p_struct_title, "Структура вебсайту Кваліфікаційного центру за розділами:", bold=True, size=12)

    sections = [
        "«Головна» (інформація про Кваліфікаційний центр, вимоги до кваліфікацій та калькулятор відповідності);",
        "«Нормативна база» (нормативно-правові акти, професійний стандарт, технічні вимоги та паспорти МТБ);",
        "«Подати заяву» (електронна форма для реєстрації та подання документів здобувачем);",
        "«Реєстр» (відкрита база даних виданих сертифікатів)."
    ]

    for item in sections:
        p_item = add_p(doc, space_after=4, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        p_item.paragraph_format.left_indent = Cm(1.25)
        add_run(p_item, "• ", bold=True, size=12)
        add_run(p_item, item, size=12)

    # Spacer
    add_p(doc, "", space_after=36)

    # 4. Signature Block
    # We will lay out the signature block using tab stops or table to keep it aligned,
    # but a simple right-aligned / two-column layout works perfectly. Let's do it clean:
    p_sig = add_p(doc, space_after=4, space_before=24)
    # Use tab stop to align name to the right
    # Left tab stop at 0, Right tab stop at 16.5cm
    p_sig.paragraph_format.tab_stops.add_tab_stop(Cm(11.0))
    add_run(p_sig, "Ректор КЗ «ЗОІППО» ЗОР\tЕ.А. ГУГНІН", bold=True, size=12)

    p_date = add_p(doc, space_after=4, space_before=18)
    p_date.paragraph_format.tab_stops.add_tab_stop(Cm(11.0))
    add_run(p_date, "«___» ____________ 2026 р.", size=12)

    # Save
    out_path = 'docs/19. про сайт ЗОІППО.docx'
    doc.save(out_path)
    print(f"Generated successfully: {out_path}")

if __name__ == '__main__':
    build_site_info_docx()
