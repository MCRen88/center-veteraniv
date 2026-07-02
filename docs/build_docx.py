import os
import json
import re
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def clean_option(line):
    line = line.strip()
    clean = line.lstrip('*- ')
    if clean.startswith('**'):
        clean = clean[2:]
    paren_idx = clean.find(')')
    if paren_idx != -1 and paren_idx < 5:
        clean = clean[paren_idx+1:].strip()
    if clean.startswith('**'):
        clean = clean[2:]
    if clean.endswith('**'):
        clean = clean[:-2]
    # Clean checkmark if present
    clean = re.sub(r'\s*✓\s*(\(Правильна відповідь\))?', '', clean).strip()
    return clean.strip()

def parse_cases(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split content into cases
    case_blocks = re.split(r'### Кейс \d+\.', content)[1:]
    cases = []
    
    for idx, block in enumerate(case_blocks):
        lines = block.strip().split('\n')
        title = lines[0].strip()
        
        situation = ""
        question = ""
        options = []
        correct_answer = ""
        justification = ""
        
        state = None # 'situation', 'question', 'options', 'justification'
        
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            
            # Check for fields
            if '* **Ситуація**:' in line or '* **Опис ситуації**:' in line:
                state = 'situation'
                situation = re.sub(r'^\*\s+\*\*Ситуація\*\*:\s*|^\*\s+\*\*Опис ситуації\*\*:\s*', '', line_str)
                continue
            elif '* **Запитання**:' in line or '* **Питання**:' in line:
                state = 'question'
                question = re.sub(r'^\*\s+\*\*Запитання\*\*:\s*|^\*\s+\*\*Питання\*\*:\s*', '', line_str)
                continue
            elif '* **Варіанти відповідей**:' in line:
                state = 'options'
                continue
            elif '* **Правильна відповідь**:' in line:
                state = 'correct'
                correct_answer = re.sub(r'^\*\s+\*\*Правильна відповідь\*\*:\s*', '', line_str).replace('*', '').strip()
                continue
            elif '* **Обґрунтування**:' in line or '* **Обґрунтування відповіді**:' in line:
                state = 'justification'
                justification = re.sub(r'^\*\s+\*\*Обґрунтування\*\*:\s*|^\*\s+\*\*Обґрунтування відповіді\*\*:\s*', '', line_str)
                continue
                
            # Accumulate content based on state
            if state == 'situation':
                situation += " " + line_str
            elif state == 'question':
                if line_str.startswith('*') or line_str.startswith('-'):
                    options.append(clean_option(line_str))
            elif state == 'justification':
                justification += " " + line_str
                
        cases.append({
            "id": idx + 1,
            "title": title,
            "situation": situation.strip(),
            "question": question.strip(),
            "options": options,
            "correct": correct_answer,
            "justification": justification.strip()
        })
        
    return cases

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell margins (padding) in dxa (1/20 of a pt)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_p(doc, text="", bold=False, italic=False, size=12, alignment=None, space_after=4, space_before=0, line_spacing=1.15):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        p.alignment = alignment
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

def add_run(p, text, bold=False, italic=False, size=12):
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def build_docx():
    print("Loading test questions...")
    with open('docs/questions.json', 'r', encoding='utf-8') as f:
        questions = json.load(f)
    print(f"Loaded {len(questions)} test questions.")

    print("Loading cases...")
    cases = parse_cases('docs/professional_cases.md')
    print(f"Loaded {len(cases)} professional cases.")

    # Initialize Document
    doc = Document()

    # Page Margins Setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # 1. TITLE PAGE (COVER)
    print("Generating Cover Page...")
    add_p(doc, "ЗАТВЕРДЖЕНО", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_p(doc, "Керівник Центру післядипломної освіти та", size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_p(doc, "сертифікацій Запорізького", size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_p(doc, "національного університету", size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_p(doc, "____________ Сергій ІЛЬЇН", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_p(doc, '"___" __________ 2026 р.', size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=40)

    # Center-aligned main title block
    add_p(doc, "КОНТРОЛЬНО-ОЦІНЮВАЛЬНІ МАТЕРІАЛИ", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_p(doc, "для проведення оцінювання результатів навчання за професійною кваліфікацією", size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_p(doc, "«Фахівець із супроводу ветеранів війни та демобілізованих осіб»", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    # Spacers and City/Year
    for _ in range(12):
        add_p(doc, "", space_after=6)
    add_p(doc, "Запоріжжя", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_p(doc, "2026", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # 2. TEST QUESTIONS (1-70)
    print("Generating Test Questions...")
    add_p(doc, "КОНТРОЛЬНО-ОЦІНЮВАЛЬНІ МАТЕРІАЛИ", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_p(doc, "(тестові завдання)", size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    add_p(doc, "Вказівка: Оберіть одну правильну відповідь відповідно до змісту запитання.", italic=True, size=11, space_after=18)

    opt_letters = ['а', 'б', 'в', 'г']
    for idx, q in enumerate(questions):
        p = add_p(doc, space_after=4, space_before=12)
        p.paragraph_format.keep_with_next = True
        add_run(p, f"{idx + 1}. ", bold=True)
        add_run(p, q['question'], bold=True)
        
        for o_idx, opt in enumerate(q['options']):
            opt_letter = opt_letters[o_idx] if o_idx < len(opt_letters) else str(o_idx + 1)
            p_opt = add_p(doc, space_after=3, space_before=0)
            p_opt.paragraph_format.left_indent = Cm(1.0)
            add_run(p_opt, f"{opt_letter}) ", italic=True)
            add_run(p_opt, opt)

    add_p(doc, "Термін виконання – 3 години", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=24, space_after=24)

    doc.add_page_break()

    # 3. PRACTICAL CASES (1-10)
    print("Generating Practical Cases...")
    add_p(doc, "КОМПЛЕКСНЕ КВАЛІФІКАЦІЙНЕ (ПРАКТИЧНЕ) ЗАВДАННЯ", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    for idx, c in enumerate(cases):
        add_p(doc, f"Практичний кейс № {idx + 1}. {c['title']}", bold=True, size=12, space_before=18, space_after=6)
        
        p_sit = add_p(doc, space_after=4)
        add_run(p_sit, "Опис ситуації: ", bold=True)
        add_run(p_sit, c['situation'])
        
        p_q = add_p(doc, space_after=4, space_before=6)
        add_run(p_q, "Питання: ", bold=True)
        add_run(p_q, c['question'], bold=True)
        
        for o_idx, opt in enumerate(c['options']):
            opt_letter = opt_letters[o_idx] if o_idx < len(opt_letters) else str(o_idx + 1)
            p_opt = add_p(doc, space_after=3, space_before=0)
            p_opt.paragraph_format.left_indent = Cm(1.0)
            add_run(p_opt, f"{opt_letter}) ", italic=True)
            add_run(p_opt, opt)

    add_p(doc, "Термін виконання – 3 години", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_before=24, space_after=24)

    doc.add_page_break()

    # 4. ANSWER KEYS & EVALUATION CRITERIA
    print("Generating Answer Keys...")
    add_p(doc, "МЕТОДИЧНІ МАТЕРІАЛИ ТА КЛЮЧІ ВІДПОВІДЕЙ", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_p(doc, "(для членів кваліфікаційної комісії)", size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # Test keys table
    add_p(doc, "1. Ключі відповідей до тестових завдань (70 питань)", bold=True, size=12, space_before=12, space_after=6)

    # 6-column table:
    # Col 0: №, Col 1: Відповідь, Col 2: №, Col 3: Відповідь, Col 4: №, Col 5: Відповідь
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Відповідь'
    hdr_cells[2].text = '№'
    hdr_cells[3].text = 'Відповідь'
    hdr_cells[4].text = '№'
    hdr_cells[5].text = 'Відповідь'
    
    # Format header
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(cell, top=100, bottom=100, left=60, right=60)
        set_cell_border(cell, 
                        top={'sz': 12, 'val': 'single', 'color': '000000'},
                        bottom={'sz': 12, 'val': 'single', 'color': '000000'},
                        left={'sz': 12, 'val': 'single', 'color': '000000'},
                        right={'sz': 12, 'val': 'single', 'color': '000000'})

    # Make keys array
    keys_list = []
    for idx, q in enumerate(questions):
        correct_idx = q['correct']
        correct_letter = opt_letters[correct_idx] if correct_idx < len(opt_letters) else 'н/д'
        keys_list.append((idx + 1, correct_letter.upper()))

    # Rows count = ceil(70 / 3) = 24 rows
    rows_count = (len(keys_list) + 2) // 3
    for r_idx in range(rows_count):
        row_cells = table.add_row().cells
        
        # Col 0 & 1
        item1 = keys_list[r_idx]
        row_cells[0].text = str(item1[0])
        row_cells[1].text = item1[1]
        
        # Col 2 & 3
        idx2 = r_idx + rows_count
        if idx2 < len(keys_list):
            item2 = keys_list[idx2]
            row_cells[2].text = str(item2[0])
            row_cells[3].text = item2[1]
        else:
            row_cells[2].text = ""
            row_cells[3].text = ""
            
        # Col 4 & 5
        idx3 = r_idx + 2 * rows_count
        if idx3 < len(keys_list):
            item3 = keys_list[idx3]
            row_cells[4].text = str(item3[0])
            row_cells[5].text = item3[1]
        else:
            row_cells[4].text = ""
            row_cells[5].text = ""
            
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_margins(cell, top=80, bottom=80, left=60, right=60)
            set_cell_border(cell, 
                            top={'sz': 4, 'val': 'single', 'color': '808080'},
                            bottom={'sz': 4, 'val': 'single', 'color': '808080'},
                            left={'sz': 4, 'val': 'single', 'color': '808080'},
                            right={'sz': 4, 'val': 'single', 'color': '808080'})

    add_p(doc, "", space_after=18)

    # Keys to Practical Cases (Rationales)
    add_p(doc, "2. Відповіді та обґрунтування до практичних кейсів (10 кейсів)", bold=True, size=12, space_before=12, space_after=6)
    
    for idx, c in enumerate(cases):
        add_p(doc, f"Практичний кейс № {idx + 1}. {c['title']}", bold=True, size=11, space_before=8, space_after=3)
        p_ans = add_p(doc, space_after=3)
        add_run(p_ans, "Правильна відповідь: ", bold=True)
        add_run(p_ans, f"{c['correct'].upper()}", bold=True)
        
        p_just = add_p(doc, space_after=4)
        p_just.paragraph_format.left_indent = Cm(0.5)
        add_run(p_just, "Обґрунтування: ", italic=True)
        add_run(p_just, c['justification'])

    # Save document
    out_path = 'docs/КОМ_Фахівець_із_супроводу.docx'
    doc.save(out_path)
    print(f"Document saved successfully to: {out_path}")

if __name__ == '__main__':
    build_docx()
