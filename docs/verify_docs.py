import docx
from docx import Document

def verify_zayava():
    print("Verifying Заява.docx...")
    doc = Document('docs/Заява.docx')
    
    # 1. Check title
    title_paragraphs = [p.text for p in doc.paragraphs[:5] if p.text.strip()]
    print("  Title paragraphs:", title_paragraphs)
    
    # 2. Check fields in the table
    tbl = doc.tables[0]
    applicant_name = tbl.rows[0].cells[1].text
    standard_info = tbl.rows[6].cells[1].text
    qualification_name = tbl.rows[7].cells[1].text
    
    print(f"  Applicant: '{applicant_name}'")
    print(f"  Standard: '{standard_info}'")
    print(f"  Qualification: '{qualification_name}'")
    
    assert "Організаційно" not in applicant_name, "Error: structural unit still in applicant name!"
    assert "10.10.2025" in standard_info and "835" in standard_info, "Error: standard not updated to 835!"
    assert "2446.2" not in qualification_name, "Error: qualification code still in name!"
    
    # 3. Check item 8 of the documents list
    found_item_8 = False
    for p in doc.paragraphs:
        if p.text.startswith("8."):
            print(f"  Item 8: '{p.text}'")
            found_item_8 = True
            assert "2446.2" not in p.text, "Error: classification code in item 8 description!"
            assert "назва професійної кваліфікації «Фахівець із супроводу" in p.text, "Error: item 8 qualification name incorrect!"
            
    assert found_item_8, "Error: item 8 of the document list not found!"
    print("Заява.docx is CORRECT!\n")

def verify_mtz():
    print("Verifying Відомості_МТЗ_ЗОІППО.docx...")
    doc = Document('docs/Відомості_МТЗ_ЗОІППО.docx')
    
    # 0. Check section orientation
    assert str(doc.sections[0].orientation) == "LANDSCAPE (1)" or doc.sections[0].page_width.cm > doc.sections[0].page_height.cm, "Error: MTZ document is not in Landscape orientation!"

    # 1. Check header paragraph 2
    header_text = doc.paragraphs[2].text
    print(f"  Header paragraph: '{header_text}'")
    assert "Організаційно-методичний центр" in header_text, "Error: structural unit name missing in MTZ header!"
    assert "Кваліфікаційний центр" not in header_text.split("Код ЄДРПОУ:")[-1], "Error: still says 'Кваліфікаційний центр' at the end!"

    # 2. Check Table 0 (theoretical tests) headers (7 columns)
    tbl0 = doc.tables[0]
    headers0 = [cell.text.strip().replace('\n', ' ') for cell in tbl0.rows[0].cells]
    print("  Table 0 headers:", headers0)
    assert len(headers0) == 7, f"Error: Table 0 should have 7 columns, found {len(headers0)}"
    assert "Назва банку тестових завдань" in headers0[1], "Error: Table 0 col 1 header mismatch!"
    assert "Кількість завдань, варіантів" in headers0[2], "Error: Table 0 col 2 header mismatch!"

    # 3. Check Table 1 (other COMs) headers (5 columns)
    tbl1 = doc.tables[1]
    headers1 = [cell.text.strip().replace('\n', ' ') for cell in tbl1.rows[0].cells]
    print("  Table 1 headers:", headers1)
    assert len(headers1) == 5, f"Error: Table 1 should have 5 columns, found {len(headers1)}"

    # 4. Check Table 2 (equipment & inventory numbers)
    tbl2 = doc.tables[2]
    print("  Table 2 (equipment) column 5 updates:")
    for row_idx in range(1, 6):
        item_name = tbl2.rows[row_idx].cells[1].text.strip().replace('\n', ' ')
        grounds = tbl2.rows[row_idx].cells[5].text.strip().replace('\n', ' ')
        print(f"    {item_name} -> {grounds}")
        assert "інв. №" in grounds or "інвентарні" in grounds or "інв. №" in grounds.lower(), f"Error: inventory number missing in row {row_idx}!"

    print("Відомості_МТЗ_ЗОІППО.docx is CORRECT!\n")

def verify_kom():
    print("Verifying КОМ_Фахівець_із_супроводу.docx...")
    doc = Document('docs/КОМ_Фахівець_із_супроводу.docx')
    
    # Check approval block on cover page
    found_rector = False
    found_university = False
    found_ilyin = False
    found_date = False
    
    for p in doc.paragraphs[:20]:
        if "Едуард ГУГНІН" in p.text:
            found_rector = True
            print(f"  Rector: '{p.text}'")
        if "Запорізького національного університету" in p.text:
            found_university = True
        if "Сергій ІЛЬЇН" in p.text:
            found_ilyin = True
        if "липня 2026" in p.text:
            found_date = True
            print(f"  Approval Date: '{p.text}'")
            
    assert found_rector, "Error: Eduard GUGNIN missing in COM approval block!"
    assert not found_university, "Error: Zaporizhzhia National University still in COM approval block!"
    assert not found_ilyin, "Error: Sergey Ilyin still in COM approval block!"
    assert found_date, "Error: Approval Date 'липня 2026' missing in COM approval block!"
    
    print("КОМ_Фахівець_із_супроводу.docx is CORRECT!\n")

def verify_zvit():
    print("Verifying Звіт_з_апробації.docx...")
    doc = Document('docs/Звіт_з_апробації.docx')
    
    # Check letterhead
    header_texts = [p.text for p in doc.paragraphs[:10] if p.text.strip()]
    print("  Header text snippet:", header_texts[:5])
    
    # Check ref and date
    found_ref = False
    for p in doc.paragraphs:
        if "№ 01-12/452" in p.text:
            found_ref = True
            print(f"  Ref/Date paragraph: '{p.text}'")
            
    assert found_ref, "Error: reference number '№ 01-12/452' missing in report!"
    
    # Check dates in content
    found_aprobation_dates = False
    found_validity_date = False
    
    for p in doc.paragraphs:
        if "02 лютого по 25 березня 2026 року" in p.text or "02.02.2026 - 25.03.2026" in p.text or "02.02.2026" in p.text:
            found_aprobation_dates = True
        if "26 березня 2026 року" in p.text:
            found_validity_date = True
            
    assert found_aprobation_dates, "Error: approbation date '02.02.2026 - 25.03.2026' incorrect or missing!"
    assert found_validity_date, "Error: validity date '26 березня 2026 року' missing!"
    
    print("Звіт_з_апробації.docx is CORRECT!\n")

def main():
    verify_zayava()
    verify_mtz()
    verify_kom()
    verify_zvit()
    print("ALL VERIFICATIONS PASSED SUCCESSFULLY!")

if __name__ == '__main__':
    main()
